import email
import imaplib
import os
import smtplib
# Добавляем необходимые подклассы - MIME-типы
import mimetypes  # Импорт класса для обработки неизвестных MIME-типов, базирующихся на расширении файла
import email
from email import encoders  # Импортируем энкодер
from email.header import decode_header
from email.mime.base import MIMEBase  # Общий тип
from email.mime.text import MIMEText  # Текст/HTML
from email.mime.image import MIMEImage  # Изображения
from email.mime.audio import MIMEAudio  # Аудио
from email.mime.multipart import MIMEMultipart  # Многокомпонентный объект
import zipfile
import webbrowser as wb
import docx


class Email:
    addr_from = 'testTest7771115@gmail.com'
    password = 'test5111777'

    def send_email(self, addr_to, msg_subj, msg_text, files):
        msg = MIMEMultipart()  # Создаем сообщение
        msg['From'] = self.addr_from  # Адресат
        msg['To'] = addr_to  # Получатель
        msg['Subject'] = msg_subj  # Тема сообщения

        body = msg_text  # Текст сообщения
        msg.attach(MIMEText(body, 'plain'))  # Добавляем в сообщение текст

        self.process_attachement(msg, files)

        # ======== Этот блок настраивается для каждого почтового провайдера отдельно ===============================================
        server = smtplib.SMTP('smtp.gmail.com', 587)  # Создаем объект SMTP
        server.starttls()  # Начинаем шифрованный обмен по TLS
        # server.set_debuglevel(True)                            # Включаем режим отладки, если не нужен - можно закомментировать
        server.login(self.addr_from, self.password)  # Получаем доступ
        server.send_message(msg)  # Отправляем сообщение
        server.quit()  # Выходим
        # ==========================================================================================================================

    def process_attachement(self, msg, files):  # Функция по обработке списка, добавляемых к сообщению файлов
        for f in files:
            if os.path.isfile(f):  # Если файл существует
                self.attach_file(msg, f)  # Добавляем файл к сообщению.
            elif os.path.exists(f):  # Если путь не файл и существует, значит - папка
                dir = os.listdir(f)  # Получаем список файлов в папке
                for file in dir:  # Перебираем все файлы и...
                    self.attach_file(msg, f + "/" + file)  # ...добавляем каждый файл к сообщению

    def attach_file(self, msg, filepath):  # Функция по добавлению конкретного файла к сообщению
        filename = os.path.basename(filepath)  # Получаем только имя файла
        ctype, encoding = mimetypes.guess_type(filepath)  # Определяем тип файла на основе его расширения
        if ctype is None or encoding is not None:  # Если тип файла не определяется
            ctype = 'application/octet-stream'  # Будем использовать общий тип
        maintype, subtype = ctype.split('/', 1)  # Получаем тип и подтип
        if maintype == 'text':  # Если текстовый файл
            with open(filepath) as fp:  # Открываем файл для чтения
                file = MIMEText(fp.read(), _subtype=subtype)  # Используем тип MIMEText
                fp.close()  # После использования файл обязательно нужно закрыть
        elif maintype == 'image':  # Если изображение
            with open(filepath, 'rb') as fp:
                file = MIMEImage(fp.read(), _subtype=subtype)
                fp.close()
        elif maintype == 'audio':  # Если аудио
            with open(filepath, 'rb') as fp:
                file = MIMEAudio(fp.read(), _subtype=subtype)
                fp.close()
        else:  # Неизвестный тип файла
            with open(filepath, 'rb') as fp:
                filename, file_extension = os.path.splitext(filepath)
                file = MIMEBase(maintype, subtype)  # Используем общий MIME-тип
                file.set_payload(fp.read())  # Добавляем содержимое общего типа (полезную нагрузку)
                fp.close()
                encoders.encode_base64(file)  # Содержимое должно кодироваться как Base64
        file.add_header('Content-Disposition', 'attachment', filename=filename)  # Добавляем заголовки
        msg.attach(file)  # Присоединяем файл к сообщению


files = ["archive_name.zip"]  # Если нужно отправить все файлы из заданной папки, нужно указать её
Email().send_email(Email.addr_from, "Тема сообщения 3", "Текст сообщения 3", files)

mail = imaplib.IMAP4_SSL('imap.gmail.com')
mail.login(Email.addr_from, Email.password)

last_mail = mail.select("inbox")[1][0]
# извлекает сообщение электронной почты по идентификатору, используя стандартный формат, указанный в RFC 822
res, msg = mail.fetch(last_mail, "(RFC822)")
for response in msg:
    if isinstance(response, tuple):
        # анализировать байтовое электронное письмо в объект сообщения
        msg = email.message_from_bytes(response[1])
        # расшифровка темы письма
        subject, encoding = decode_header(msg["Subject"])[0]
        if isinstance(subject, bytes):
            # если это байты, декодировать в str
            subject = subject.decode(encoding)
        # расшифровать отправителя электронной почты
        From, encoding = decode_header(msg.get("From"))[0]
        if isinstance(From, bytes):
            From = From.decode(encoding)
        print("From:", From)
        print("Subject:", subject)
        # если сообщение электронной почты составное
        if msg.is_multipart():
            # перебираем части письма
            for part in msg.walk():
                # извлекаем тип содержимого электронной почты
                content_type = part.get_content_type()
                content_disposition = str(part.get("Content-Disposition"))
                try:
                    # получаем тело письма
                    body = part.get_payload(decode=True).decode()
                except:
                    pass
                if content_type == "text/plain" and "attachment" not in content_disposition:
                    # вывести text/plain электронные письма и пропускать вложения
                    print(body)
                elif "attachment" in content_disposition:
                    # скачать вложение
                    filename = part.get_filename()
                    if filename:
                        if not os.path.isdir(subject):
                            # создайте папку для этого электронного письма (названную в соответствии с темой)
                            os.mkdir(subject)
                        filepath = os.path.join(subject, filename)
                        # скачать вложение и сохранить его
                        open(filepath, "wb").write(part.get_payload(decode=True))
                        # archive = zipfile.ZipFile('archive_name.zip', 'r')
                        # txtdata = archive.read("Романов Роман Максимович/Задания/Задание1.docx")
                        with zipfile.ZipFile(subject + '/' + filename) as myzip:
                            myzip.extract('Романов Роман Максимович/Задания/Задание1.docx', '.')
                            myzip.extract('Романов Роман Максимович/Задания/Задание2.docx', '.')
                        doc1 = docx.Document("Романов Роман Максимович/Задания/Задание1.docx")
                        doc2 = docx.Document("Романов Роман Максимович/Задания/Задание2.docx")
                        wb.open('https://yandex.ru/search/?text=' + doc1.paragraphs[0].text)
                        wb.open('https://yandex.ru/search/?text=' + doc2.paragraphs[0].text)
        else:
            # извлекать тип содержимого электронной почты
            content_type = msg.get_content_type()
            # получить тело письма
            body = msg.get_payload(decode=True).decode()
            if content_type == "text/plain":
                # печатать только текстовые части электронного письма
                print(body)
        if content_type == "text/html":
            # если это HTML, создайте новый HTML-файл и откройте его в браузере.
            if not os.path.isdir(subject):
                # создайте папку для этого электронного письма (названную в соответствии с темой)
                os.mkdir(subject)
            filename = "index.html"
            filepath = os.path.join(subject, filename)
            # напишите файл
            open(filepath, "w").write(body)
            # открыть в браузере по умолчанию
            wb.open(filepath)
        print("=" * 100)
# закройте соединение и выйдите из системы
mail.close()
mail.logout()
