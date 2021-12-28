import getpass
import os
import shutil

from openpyxl import load_workbook
import docx
import zipfile

# 1) Парсинг excel, занести в список
# 2) Папка на рабочем столе с номерами группы
# 3) Для каждого студента создав папку с его ФИО

wb = load_workbook('Список группы.xlsx')
sheet_ranges = wb['Лист1']
column_b = sheet_ranges['B']

doc1 = docx.Document("П 2/Задание1.docx")
all_paras1 = doc1.paragraphs
doc2 = docx.Document("П 2/Задание2.docx")
all_paras2 = doc2.paragraphs
all_exercise = {}
i = 1
for para1 in all_paras1:
    if sheet_ranges.max_row >= i:
        all_exercise[i] = [para1.text]
        i += 1
    else:
        break
i = 1
for para2 in all_paras2:
    if sheet_ranges.max_row >= i:
        all_exercise[i].append(para2.text)
        i += 1
    else:
        break
students = []
# Или просто как итератор перебрать столбец:
for cell in column_b:
    students.append(cell.value)
dirFolder = os.getenv("SystemDrive") + '\\Users\\' + getpass.getuser() + '\\Desktop\\407'
i = 1
for student in students:
    pathExercise = dirFolder + '\\' + student + '\\' + 'Задания'
    os.makedirs(pathExercise)
    docExercise1 = docx.Document()
    docExercise1.add_paragraph(all_exercise[i][0])
    docExercise1.save(pathExercise + '\\' + 'Задание1.docx')
    docExercise2 = docx.Document()
    docExercise2.add_paragraph(all_exercise[i][1])
    docExercise2.save(pathExercise + '\\' + 'Задание2.docx')
    i += 1
if os.path.exists(dirFolder):
    archive_name = os.path.expanduser(os.path.join(dirFolder, 'myarchive'))
    root_dir = os.path.expanduser(os.path.join('~', dirFolder))
    shutil.make_archive('archive_name', 'zip', root_dir)
else:
    print('Папка или файл не найден')
