import getpass
import os
import shutil

from openpyxl import load_workbook
import docx

# 1) Парсинг excel, занести в список
# 2) Папка на рабочем столе с номерами группы
# 3) Для каждого студента создав папку с его ФИО

# берёт текущий путь каталога
current_dir = os.path.dirname(__file__)
# к current_dir доавляем путь
file_path = os.path.join(current_dir, "Список группы.xlsx")
print(file_path)

wb = load_workbook(file_path)
sheet_ranges = wb['Лист1']  # берёт первую страницу в файле excel
column_b = sheet_ranges['B']  # берёт столбец B

# к current_dir доавляем путь
doc1_path = os.path.join(current_dir, "Задание1.docx")
doc1 = docx.Document(doc1_path)
all_paras1 = doc1.paragraphs  # берёт все параграфы(строки) файла
doc2_path = os.path.join(current_dir, "Задание2.docx")
doc2 = docx.Document(doc2_path)
all_paras2 = doc2.paragraphs
# словарь с заданиями. Ключ - номер студента, значение - его два задания из двух файлов
all_exercise = {}
i = 1
for para1 in all_paras1:  # перебор каждого параграфа(строк)
    # sheet_ranges.max_row - подсчёт максимального количества строк
    if sheet_ranges.max_row >= i:
        all_exercise[i] = [para1.text]
        i += 1
    else:
        break
i = 1
for para2 in all_paras2:
    if sheet_ranges.max_row >= i:
        # добавление в список текст параграфа
        all_exercise[i].append(para2.text)
        i += 1
    else:
        break
students = []
for cell in column_b:  # перебор строк столбца B
    # добавление в список студента
    students.append(cell.value)
dirFolder = os.getenv("SystemDrive") + '\\Users\\' + \
    getpass.getuser() + '\\Desktop\\407'
i = 1
for student in students:
    pathExercise = dirFolder + '\\' + student + '\\' + 'Задания'
    os.makedirs(pathExercise) # создание пути
    docExercise1 = docx.Document() # создание word файла
    docExercise1.add_paragraph(all_exercise[i][0]) # добавляем параграф с заданием
    docExercise1.save(pathExercise + '\\' + 'Задание1.docx') # сохранение файла
    docExercise2 = docx.Document()
    docExercise2.add_paragraph(all_exercise[i][1])
    docExercise2.save(pathExercise + '\\' + 'Задание2.docx')
    i += 1
# проверка на существование пути
if os.path.exists(dirFolder):
    # возвращает аргумент с начальным компонентом пути '~' или '~user', 
    # замененным домашним каталогом этого пользователя. 
    # Если расширение пути завершается неудачно или путь не начинается с тильды '~', то путь возвращается без изменений.
    root_dir = os.path.expanduser(os.path.join('~', dirFolder))
    shutil.make_archive('archive_name', 'zip', root_dir)
else:
    print('Папка или файл не найден')
